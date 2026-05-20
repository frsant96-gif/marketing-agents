from docx import Document
from docx.shared import Pt, RGBColor, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re, os

# ── helpers ──────────────────────────────────────────────────────────────────

def add_hyperlink(para, text, url):
    part = para.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hl = OxmlElement("w:hyperlink"); hl.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    st = OxmlElement("w:rStyle"); st.set(qn("w:val"), "Hyperlink")
    c = OxmlElement("w:color"); c.set(qn("w:val"), "0563C1")
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single")
    rPr.append(st); rPr.append(c); rPr.append(u); r.append(rPr)
    t = OxmlElement("w:t"); t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t); hl.append(r); para._p.append(hl)

def inline_links(para, text):
    tokens = re.split(r'(\*\*[^*]+\*\*|\[[^\]]+\]\([^\)]+\))', text)
    for tok in tokens:
        bm = re.match(r'^\*\*([^*]+)\*\*$', tok)
        lm = re.match(r'^\[([^\]]+)\]\(([^\)]+)\)$', tok)
        if bm:
            para.add_run(bm.group(1)).bold = True
        elif lm:
            add_hyperlink(para, lm.group(1), lm.group(2))
        elif tok:
            para.add_run(tok)

def inline(para, text):
    clean = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    for part in re.split(r"(\*\*[^*]+\*\*)", clean):
        if part.startswith("**") and part.endswith("**"):
            para.add_run(part[2:-2]).bold = True
        else:
            para.add_run(part)

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br"); br.set(qn("w:type"), "page")
    run._r.append(br)

def add_code_block(doc, code_text):
    for line in code_text.split("\n"):
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Pt(18)
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        # fundo cinza claro
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "F2F2F2")
        pPr.append(shd)

# ── parser de markdown ────────────────────────────────────────────────────────

AZUL = RGBColor(0x00, 0x2B, 0x5C)
MEIO = RGBColor(0x00, 0x6B, 0xA8)

def heading(doc, level, text):
    para = doc.add_paragraph()
    para.style = doc.styles[f"Heading {level}"]
    r = para.add_run(text); r.bold = True
    r.font.color.rgb = AZUL if level <= 2 else MEIO
    r.font.size = Pt(20 if level == 1 else 14 if level == 2 else 12)

def parse_md(doc, lines, use_links=False):
    in_fontes = False
    in_code = False
    code_buf = []
    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n")

        # code block
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True; code_buf = []; i += 1; continue
            else:
                add_code_block(doc, "\n".join(code_buf))
                doc.add_paragraph()
                in_code = False; code_buf = []; i += 1; continue
        if in_code:
            code_buf.append(line); i += 1; continue

        if line.strip() == "## Fontes":
            in_fontes = True
            heading(doc, 2, "Fontes"); i += 1; continue

        if line.startswith("# ") and not line.startswith("## "):
            in_fontes = False
            heading(doc, 1, line[2:]); i += 1; continue
        if line.startswith("## "):
            in_fontes = False
            heading(doc, 2, line[3:]); i += 1; continue
        if line.startswith("### "):
            heading(doc, 3, line[4:]); i += 1; continue
        if line.strip() in ("---", ""):
            i += 1; continue

        if line.startswith(("- ", "→ ")):
            p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(4)
            if in_fontes or use_links:
                inline_links(p, line[2:])
            else:
                inline(p, line[2:])
            i += 1; continue

        if re.match(r"^\d+\.\s", line):
            p = doc.add_paragraph(style="List Number"); p.paragraph_format.space_after = Pt(4)
            inline(p, re.sub(r"^\d+\.\s", "", line)); i += 1; continue

        if re.match(r"^\*\*[^*]+\*\*$", line):
            p = doc.add_paragraph(); r = p.add_run(line[2:-2]); r.bold = True
            r.font.size = Pt(12); r.font.color.rgb = MEIO
            p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
            i += 1; continue

        if line.startswith("|"):
            tbl_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                tbl_lines.append(lines[i].rstrip("\n")); i += 1
            rows = [l for l in tbl_lines if not re.match(r"^\|[-| :]+\|$", l)]
            if rows:
                nc = len(rows[0].split("|")) - 2
                tbl = doc.add_table(rows=len(rows), cols=nc); tbl.style = "Table Grid"
                for ri, rl in enumerate(rows):
                    cells = [c.strip() for c in rl.split("|")[1:-1]]
                    for ci in range(min(len(cells), nc)):
                        # limpa backslash em pipes escapados
                        cell_text = cells[ci].replace("\\|", "|")
                        tbl.rows[ri].cells[ci].text = cell_text
                        if ri == 0:
                            for pp in tbl.rows[ri].cells[ci].paragraphs:
                                for rr in pp.runs: rr.bold = True
            doc.add_paragraph(); continue

        if line.startswith("**[") and line.endswith("]**"):
            p = doc.add_paragraph(); r = p.add_run(line[3:-3]); r.bold = True
            r.font.size = Pt(12); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(10)
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "002B5C")
            p._p.get_or_add_pPr().append(shd); i += 1; continue

        p = doc.add_paragraph(style="Normal"); p.paragraph_format.space_after = Pt(6)
        inline(p, line); i += 1

# ── gerador principal ─────────────────────────────────────────────────────────

def gerar(artigo_path, meta_path, output_path):
    with open(artigo_path, encoding="utf-8") as f:
        artigo_lines = f.readlines()
    with open(meta_path, encoding="utf-8") as f:
        meta_lines = f.readlines()

    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Pt(72)
        s.left_margin = s.right_margin = Pt(90)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    parse_md(doc, artigo_lines)
    add_page_break(doc)
    parse_md(doc, meta_lines)

    doc.save(output_path)
    print(f"OK: {output_path}")

BASE = os.path.dirname(os.path.abspath(__file__))
for slug in [
    "sap-business-ai-platform",
    "sap-autonomous-suite",
    "sap-knowledge-graph",
    "sap-joule-work",
    "sap-anthropic-parceria",
    "sap-bdc-knowledge-core",
]:
    pasta = os.path.join(BASE, slug)
    gerar(
        os.path.join(pasta, "artigo.md"),
        os.path.join(pasta, "meta.md"),
        os.path.join(pasta, "artigo.docx"),
    )
