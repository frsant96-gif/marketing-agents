import re
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = os.path.dirname(os.path.abspath(__file__))

FILES = ["anuncios.md", "posts.md", "emails.md", "landing-page-whatsapp.md"]

BLUE = RGBColor(0x00, 0x6A, 0xFF)
DARK = RGBColor(0x0A, 0x0E, 0x19)


def add_inline_runs(paragraph, text):
    # handle **bold** segments
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)


def convert(md_path, docx_path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    with open(md_path, encoding="utf-8") as f:
        lines = f.readlines()

    for raw_line in lines:
        line = raw_line.rstrip("\n")

        if not line.strip():
            continue

        if line.startswith("# "):
            h = doc.add_heading(level=1)
            add_inline_runs(h, line[2:].strip())
            h.runs[0].font.color.rgb = DARK if h.runs else None
            continue

        if line.startswith("## "):
            h = doc.add_heading(level=2)
            add_inline_runs(h, line[3:].strip())
            continue

        if line.startswith("### "):
            h = doc.add_heading(level=3)
            add_inline_runs(h, line[4:].strip())
            continue

        if line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            run = p.add_run(line[2:].strip())
            run.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            continue

        if line.startswith("---"):
            doc.add_paragraph("_" * 40)
            continue

        if line.startswith("| "):
            # skip raw table separator rows, render table rows as plain text lines
            if re.match(r"^\|[\s:\-|]+\|$", line):
                continue
            cells = [c.strip() for c in line.strip("|").split("|")]
            p = doc.add_paragraph()
            p.add_run(" | ".join(cells)).font.size = Pt(10)
            continue

        if line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, line[2:].strip())
            continue

        if re.match(r"^\d+\.\s", line):
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, re.sub(r"^\d+\.\s", "", line))
            continue

        p = doc.add_paragraph()
        add_inline_runs(p, line)

    doc.save(docx_path)


if __name__ == "__main__":
    for name in FILES:
        md_path = os.path.join(BASE, name)
        docx_path = os.path.join(BASE, name.replace(".md", ".docx"))
        convert(md_path, docx_path)
        print(f"Gerado: {docx_path}")
